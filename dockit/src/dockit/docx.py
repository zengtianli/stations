"""Word document processing — text formatting, extraction, style cleanup, format checking.

Pure logic module: bytes in, bytes out. No file paths, no console output.

Usage:
    from dockit.docx import format_text, extract_text, cleanup_styles

    result = format_text(doc_bytes)        # Fix quotes/punctuation/units
    md = extract_text(doc_bytes)           # Extract as Markdown
    result = cleanup_styles(doc_bytes)     # Remove unused styles

    snap = check_format(doc_bytes)         # Extract format snapshot
    report = format_report(snap)           # Generate Markdown report
    cmp = compare_format(before, after)    # Compare two documents
"""

import copy
import hashlib
import json
import logging
import re
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from dockit.text import fix_all


@dataclass
class FormatResult:
    """Result of a document formatting operation."""

    data: bytes
    stats: dict[str, int] = field(default_factory=dict)


# -- Internal helpers ----------------------------------------------------------

QUOTE_CHARS = {"\u201c", "\u201d"}


def _set_run_font(run_element, font_name: str):
    """Set font for a run element (ascii + hAnsi + eastAsia)."""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hint"), "eastAsia")


def _split_run_at_quotes(run):
    """Split a run at quote characters, returning [(text, is_quote), ...] or None."""
    text = run.text
    if not text or not any(c in QUOTE_CHARS for c in text):
        return None

    segments = []
    buf = []
    for c in text:
        if c in QUOTE_CHARS:
            if buf:
                segments.append(("".join(buf), False))
                buf = []
            segments.append((c, True))
        else:
            buf.append(c)
    if buf:
        segments.append(("".join(buf), False))
    return segments


def _apply_quote_split(run, segments, quote_font: str):
    """Split run into multiple runs: quote chars get the specified font."""
    parent = run._element.getparent()

    # First segment reuses the original run
    first_text, first_is_quote = segments[0]
    run.text = first_text
    if first_is_quote:
        _set_run_font(run._element, quote_font)

    # Subsequent segments: deep-copy original run, insert after
    insert_after = run._element
    for seg_text, is_quote in segments[1:]:
        new_r = copy.deepcopy(run._element)
        # Clear old text elements from the copy
        for t_elem in new_r.findall(qn("w:t")):
            new_r.remove(t_elem)
        t_elem = OxmlElement("w:t")
        t_elem.text = seg_text
        t_elem.set(qn("xml:space"), "preserve")
        new_r.append(t_elem)

        if is_quote:
            _set_run_font(new_r, quote_font)
        else:
            # Non-quote segments: restore original font
            rPr = new_r.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                orig_rPr = run._element.find(qn("w:rPr"))
                orig_rFonts = orig_rPr.find(qn("w:rFonts")) if orig_rPr is not None else None
                if rFonts is not None and orig_rFonts is not None:
                    rPr.replace(rFonts, copy.deepcopy(orig_rFonts))
                elif rFonts is not None and orig_rFonts is None:
                    rPr.remove(rFonts)

        parent.insert(list(parent).index(insert_after) + 1, new_r)
        insert_after = new_r


def _process_paragraph(paragraph, stats: dict, quote_counter: int, quote_font: str) -> int:
    """Process a single paragraph: fix text and split quotes into separate runs.

    Returns updated quote_counter.
    """
    if not paragraph.runs:
        return quote_counter

    original_runs = list(paragraph.runs)

    for run in original_runs:
        if not run.text:
            continue

        original_text = run.text
        fixed_text, fix_stats, quote_counter = fix_all(original_text, quote_counter)

        stats["quotes"] += fix_stats["quotes"]
        stats["punctuation"] += fix_stats["punctuation"]
        stats["units"] += fix_stats["units"]

        if fixed_text != original_text:
            run.text = fixed_text

        # Split quotes into separate runs with specified font
        segments = _split_run_at_quotes(run)
        if segments:
            _apply_quote_split(run, segments, quote_font)

    return quote_counter


def _process_table(table, stats: dict, quote_counter: int, quote_font: str) -> int:
    """Process all paragraphs in a table. Returns updated quote_counter."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)
    return quote_counter


# -- Public API ----------------------------------------------------------------


def format_text(
    doc_bytes: bytes,
    *,
    fix_quotes: bool = True,
    fix_punctuation: bool = True,
    fix_units: bool = True,
    quote_font: str = "\u5b8b\u4f53",
    process_headers_footers: bool = True,
    strip_headers_footers: bool = False,
) -> FormatResult:
    """Format text in a Word document.

    Fixes quotes (with font splitting), punctuation, and unit symbols
    across all paragraphs, tables, headers, and footers.

    Args:
        doc_bytes: Raw bytes of a .docx file.
        fix_quotes: Whether to fix quote characters.
        fix_punctuation: Whether to convert English punctuation to Chinese.
        fix_units: Whether to convert Chinese unit names to symbols.
        quote_font: Font name for quote characters (default: Song Ti).
        process_headers_footers: Whether to process headers/footers.
        strip_headers_footers: If True, remove all headers/footers entirely.

    Returns:
        FormatResult with processed document bytes and replacement stats.
    """
    doc = Document(BytesIO(doc_bytes))
    stats = {"quotes": 0, "punctuation": 0, "units": 0}
    quote_counter = 0

    # Process body paragraphs
    for paragraph in doc.paragraphs:
        quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)

    # Process tables
    for table in doc.tables:
        quote_counter = _process_table(table, stats, quote_counter, quote_font)

    # Process headers/footers
    if strip_headers_footers:
        for section in doc.sections:
            sectPr = section._sectPr
            for ref in sectPr.findall(qn("w:headerReference")):
                sectPr.remove(ref)
            for ref in sectPr.findall(qn("w:footerReference")):
                sectPr.remove(ref)
    elif process_headers_footers:
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)
                for table in section.header.tables:
                    quote_counter = _process_table(table, stats, quote_counter, quote_font)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    quote_counter = _process_paragraph(paragraph, stats, quote_counter, quote_font)
                for table in section.footer.tables:
                    quote_counter = _process_table(table, stats, quote_counter, quote_font)

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return FormatResult(data=buf.getvalue(), stats=stats)


# == Text extraction ===========================================================

# W namespace for raw XML access
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _wqn(tag: str) -> str:
    """Quick namespace helper for w: prefix."""
    return f"{{{_W_NS}}}{tag}"


def _extract_paragraphs_raw(doc_bytes: bytes) -> list[dict]:
    """Extract paragraphs via raw XML for maximum fidelity."""
    paragraphs = []
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        doc_xml = zf.read("word/document.xml")
        # Build style ID → name map
        styles_map = {}
        if "word/styles.xml" in zf.namelist():
            styles_xml = zf.read("word/styles.xml")
            stree = etree.fromstring(styles_xml)
            for s in stree.iter(_wqn("style")):
                sid = s.get(_wqn("styleId"), "")
                name_elem = s.find(_wqn("name"))
                name = name_elem.get(_wqn("val"), sid) if name_elem is not None else sid
                styles_map[sid] = name

        tree = etree.fromstring(doc_xml)
        body = tree.find(_wqn("body"))
        if body is None:
            return paragraphs

        for para in body.iter(_wqn("p")):
            ppr = para.find(_wqn("pPr"))
            style_id = ""
            outline_level = -1
            if ppr is not None:
                ps = ppr.find(_wqn("pStyle"))
                if ps is not None:
                    style_id = ps.get(_wqn("val"), "")
                ol = ppr.find(_wqn("outlineLvl"))
                if ol is not None:
                    outline_level = int(ol.get(_wqn("val"), "-1"))

            style_name = styles_map.get(style_id, style_id)

            level = -1
            heading_match = re.match(r"[Hh]eading\s*(\d+)", style_name)
            if heading_match:
                level = int(heading_match.group(1))
            elif outline_level >= 0:
                level = outline_level + 1

            texts = []
            for t in para.iter(_wqn("t")):
                if t.text:
                    texts.append(t.text)
            text = "".join(texts).strip()

            if text:
                paragraphs.append({"style": style_name, "text": text, "level": level})

    return paragraphs


def _paragraphs_to_markdown(paragraphs: list[dict]) -> str:
    """Convert paragraph list to Markdown."""
    lines = []
    for p in paragraphs:
        style = p["style"].lower()
        text = p["text"]
        level = p["level"]

        if level >= 1:
            lines.append(f"\n{'#' * level} {text}\n")
        elif "表" in style or "图" in style or "caption" in style:
            lines.append(f"\n**[{text}]**\n")
        elif "题目" in style or "title" in style:
            lines.append(f"\n**{text}**\n")
        else:
            lines.append(f"\n{text}\n")

    return "\n".join(lines).strip() + "\n"


def extract_text(doc_bytes: bytes) -> str:
    """Extract text from a Word document as Markdown.

    Heading styles are converted to ``#`` headings. Table/figure captions
    are bolded. Custom styles (e.g. ZDWP) are detected via outline level.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        Markdown-formatted text.
    """
    paragraphs = _extract_paragraphs_raw(doc_bytes)
    return _paragraphs_to_markdown(paragraphs)


def extract_paragraphs(doc_bytes: bytes) -> list[dict]:
    """Extract structured paragraph data from a Word document.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        List of dicts, each with keys: style, text, level (-1 if not a heading).
    """
    return _extract_paragraphs_raw(doc_bytes)


def extract_chapters(doc_bytes: bytes) -> list[dict]:
    """Extract and split document by top-level headings.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        List of chapter dicts, each with: title, paragraphs, markdown.
    """
    paragraphs = _extract_paragraphs_raw(doc_bytes)
    chapters = []
    current: dict = {"title": "", "paragraphs": []}

    for p in paragraphs:
        if p["level"] == 1:
            if current["paragraphs"]:
                current["markdown"] = _paragraphs_to_markdown(current["paragraphs"])
                chapters.append(current)
            current = {"title": p["text"], "paragraphs": [p]}
        else:
            current["paragraphs"].append(p)

    if current["paragraphs"]:
        current["markdown"] = _paragraphs_to_markdown(current["paragraphs"])
        chapters.append(current)

    return chapters


# == Style cleanup =============================================================


@dataclass
class CleanupResult:
    """Result of style cleanup operation."""

    data: bytes
    log: list[str] = field(default_factory=list)


# Built-in style IDs that should never be deleted
_BUILTIN_KEEP = {"a", "a0"}


def _get_style_map(styles_tree) -> dict:
    """Extract styleId -> {name, type, basedOn, elem} from styles.xml."""
    result = {}
    for s in styles_tree.findall(f".//{_wqn('style')}"):
        sid = s.get(_wqn("styleId"), "")
        stype = s.get(_wqn("type"), "")
        ne = s.find(_wqn("name"))
        name = ne.get(_wqn("val"), sid) if ne is not None else sid
        based = s.find(_wqn("basedOn"))
        base_id = based.get(_wqn("val"), "") if based is not None else ""
        result[sid] = {"name": name, "type": stype, "basedOn": base_id, "elem": s}
    return result


def _get_used_style_ids(doc_tree) -> set:
    """Find all style IDs referenced in document.xml."""
    used = set()
    for tag in ("pStyle", "rStyle", "tblStyle"):
        for elem in doc_tree.iter(_wqn(tag)):
            used.add(elem.get(_wqn("val"), ""))
    used.discard("")
    return used


def _get_needed_ids(style_map: dict, used_ids: set) -> set:
    """Recursively find all basedOn dependencies."""
    needed = set(used_ids)
    queue = list(used_ids)
    while queue:
        sid = queue.pop()
        info = style_map.get(sid)
        if info and info["basedOn"] and info["basedOn"] not in needed:
            needed.add(info["basedOn"])
            queue.append(info["basedOn"])
    return needed


def cleanup_styles(
    doc_bytes: bytes,
    *,
    renames: dict[str, str] | None = None,
    merges: dict[str, str] | None = None,
    delete_unused: bool = True,
) -> CleanupResult:
    """Clean up styles in a Word document.

    Operations (applied in order):
    1. Merge styles: reassign paragraphs from one style to another.
    2. Rename styles: change display names (safe, does not change IDs).
    3. Delete unused: remove style definitions not referenced by any paragraph.

    Args:
        doc_bytes: Raw bytes of a .docx file.
        renames: ``{old_display_name: new_display_name}`` mapping.
        merges: ``{from_style_id: to_style_id}`` mapping.
        delete_unused: Whether to delete unreferenced style definitions.

    Returns:
        CleanupResult with modified document bytes and operation log.
    """
    # Load all parts from the ZIP
    files = {}
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        for info in zf.infolist():
            files[info.filename] = zf.read(info.filename)

    styles_tree = etree.fromstring(files["word/styles.xml"])
    doc_tree = etree.fromstring(files["word/document.xml"])
    style_map = _get_style_map(styles_tree)
    log: list[str] = []

    # 1. Merge styles
    if merges:
        for from_id, to_id in merges.items():
            if from_id not in style_map:
                log.append(f"Merge skipped: source '{from_id}' not found")
                continue
            if to_id not in style_map:
                log.append(f"Merge skipped: target '{to_id}' not found")
                continue
            count = 0
            for ps in doc_tree.iter(_wqn("pStyle")):
                if ps.get(_wqn("val")) == from_id:
                    ps.set(_wqn("val"), to_id)
                    count += 1
            from_name = style_map[from_id]["name"]
            to_name = style_map[to_id]["name"]
            log.append(f"Merged: {from_name} -> {to_name} ({count} paragraphs)")

    # 2. Rename styles
    if renames:
        name_to_id = {info["name"]: sid for sid, info in style_map.items()}
        for old_name, new_name in renames.items():
            sid = name_to_id.get(old_name)
            if not sid:
                log.append(f"Rename skipped: '{old_name}' not found")
                continue
            ne = style_map[sid]["elem"].find(_wqn("name"))
            if ne is not None:
                ne.set(_wqn("val"), new_name)
            log.append(f"Renamed: {old_name} -> {new_name}")

    # 3. Delete unused styles
    if delete_unused:
        used_ids = _get_used_style_ids(doc_tree)
        needed_ids = _get_needed_ids(style_map, used_ids)
        needed_ids |= _BUILTIN_KEEP
        for sid in style_map:
            if sid.startswith("TOC") or sid in ("a", "a0"):
                needed_ids.add(sid)

        deleted = []
        for sid, info in style_map.items():
            if sid not in needed_ids:
                deleted.append((sid, info["name"]))
                for s in styles_tree.findall(f".//{_wqn('style')}"):
                    if s.get(_wqn("styleId")) == sid:
                        s.getparent().remove(s)
                        break

        if deleted:
            log.append(f"Deleted {len(deleted)} unused styles")

    # Save back to ZIP
    files["word/styles.xml"] = etree.tostring(styles_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)

    return CleanupResult(data=buf.getvalue(), log=log)


# == Image caption styling =====================================================

_log = logging.getLogger(__name__)


@dataclass
class CaptionResult:
    """Result of applying image caption styles."""

    data: bytes
    images_styled: int
    captions_styled: int


def _find_style_fuzzy(doc, style_name: str) -> str | None:
    """Fuzzy-find a paragraph style in the document.

    Priority: exact match > ignore-whitespace match > substring match.
    Returns the actual style name or None.
    """
    search_normalized = style_name.replace(" ", "").replace("\u3000", "")

    exact_match = None
    space_match = None
    partial_matches: list[str] = []

    for style in doc.styles:
        if not style.name or style.type != 1:  # paragraph styles only
            continue

        if style.name == style_name:
            exact_match = style.name
            break

        style_normalized = style.name.replace(" ", "").replace("\u3000", "")
        if style_normalized == search_normalized:
            space_match = style.name

        if search_normalized in style_normalized or style_normalized in search_normalized:
            if abs(len(style_normalized) - len(search_normalized)) <= 5:
                partial_matches.append(style.name)

    if exact_match:
        return exact_match
    if space_match:
        return space_match
    if partial_matches:
        return partial_matches[0]
    return None


def _has_image(paragraph) -> bool:
    """Return True if the paragraph contains an image (w:drawing or w:pict)."""
    for run in paragraph.runs:
        for child in run._element:
            if child.tag == qn("w:drawing"):
                return True
            if child.tag == qn("w:pict"):
                return True
    return False


def _is_in_table(paragraph) -> bool:
    """Return True if the paragraph is inside a table cell."""
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag == qn("w:tc"):
            return True
        parent = parent.getparent()
    return False


def _add_blank_line_after(doc, paragraph) -> None:
    """Insert a blank paragraph after *paragraph*, unless one already exists."""
    body = doc.element.body
    para_element = paragraph._element
    para_index = list(body).index(para_element)

    # Skip if the next element is already an empty paragraph
    if para_index + 1 < len(body):
        next_element = body[para_index + 1]
        if next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph as _Para

            next_para = _Para(next_element, doc)
            if not next_para.text.strip():
                return

    new_para = OxmlElement("w:p")
    body.insert(para_index + 1, new_para)


def add_captions(
    doc_bytes: bytes,
    *,
    style_name: str = "ZDWP图名",
    add_blank_after: bool = True,
) -> CaptionResult:
    """Apply a style to image paragraphs and their captions.

    Finds paragraphs containing images (not in tables), applies the specified
    style, and optionally applies it to the next paragraph (caption) and adds
    a blank line after.

    Args:
        doc_bytes: Raw bytes of a .docx file.
        style_name: Paragraph style name to apply (fuzzy matched).
        add_blank_after: Whether to add a blank paragraph after captions.

    Returns:
        CaptionResult with processed document bytes and counts.
    """
    doc = Document(BytesIO(doc_bytes))

    # Resolve the style name via fuzzy matching
    actual_style = _find_style_fuzzy(doc, style_name)
    if actual_style is None:
        _log.warning(
            "Style '%s' not found in document; falling back to 'Normal'",
            style_name,
        )
        actual_style = "Normal"
    elif actual_style != style_name:
        _log.info(
            "Fuzzy-matched style '%s' -> '%s'", style_name, actual_style
        )

    images_styled = 0
    captions_styled = 0
    paragraphs = doc.paragraphs

    for i, paragraph in enumerate(paragraphs):
        if _is_in_table(paragraph):
            continue

        if _has_image(paragraph):
            paragraph.style = actual_style
            images_styled += 1

            # Style the next paragraph (caption) if it exists and has text
            if i + 1 < len(paragraphs):
                next_paragraph = paragraphs[i + 1]
                if not _is_in_table(next_paragraph) and next_paragraph.text.strip():
                    next_paragraph.style = actual_style
                    captions_styled += 1

                    if add_blank_after:
                        _add_blank_line_after(doc, next_paragraph)

    buf = BytesIO()
    doc.save(buf)
    return CaptionResult(
        data=buf.getvalue(),
        images_styled=images_styled,
        captions_styled=captions_styled,
    )


# == Track changes & review ====================================================


@dataclass
class ReviewResult:
    """Result of applying review rules to a document."""

    data: bytes
    count: int  # number of replacements made


# -- Namespace helpers (multi-prefix, independent of python-docx qn) -----------

_NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}
_R_NS = _NSMAP["r"]
_REL_COMMENTS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)


def _qn(tag: str) -> str:
    """Resolve a prefixed tag like ``'w:t'`` to ``'{namespace}t'``.

    Supports w: and r: prefixes.  For w:-only usage the module-level
    ``_wqn`` helper is equivalent but this one is needed for comments /
    relationships that use the r: prefix.
    """
    prefix, local = tag.split(":")
    return f"{{{_NSMAP[prefix]}}}{local}"


# -- read_changes helpers ------------------------------------------------------


def _tc_extract_text(node) -> str:
    """Extract all ``<w:t>`` text from an XML node."""
    return "".join(t.text for t in node.iter(_qn("w:t")) if t.text)


def _tc_extract_del_text(node) -> str:
    """Extract ``<w:delText>`` (or ``<w:t>``) text from a deletion node."""
    parts = [t.text for t in node.iter(_qn("w:delText")) if t.text]
    if not parts:
        parts = [t.text for t in node.iter(_qn("w:t")) if t.text]
    return "".join(parts)


# -- Public API: read_changes -------------------------------------------------


def read_changes(doc_bytes: bytes) -> dict:
    """Read track changes and comments from a Word document.

    Returns:
        Dict with ``'changes'`` list and ``'comments'`` list.
        Each change has: type (``'insert'``/``'delete'``), author, date, text.
        Each comment has: id, author, date, text.
    """
    changes: list[dict] = []
    comments: list[dict] = []

    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        doc_xml = zf.read("word/document.xml")
        tree = etree.fromstring(doc_xml)

        for ins in tree.iter(_qn("w:ins")):
            author = ins.get(_qn("w:author"), "")
            date = ins.get(_qn("w:date"), "")
            text = _tc_extract_text(ins)
            if text.strip():
                changes.append(
                    {"type": "insert", "author": author, "date": date, "text": text}
                )

        for dl in tree.iter(_qn("w:del")):
            author = dl.get(_qn("w:author"), "")
            date = dl.get(_qn("w:date"), "")
            text = _tc_extract_del_text(dl)
            if text.strip():
                changes.append(
                    {"type": "delete", "author": author, "date": date, "text": text}
                )

        if "word/comments.xml" in zf.namelist():
            comments_xml = zf.read("word/comments.xml")
            ctree = etree.fromstring(comments_xml)
            for comment in ctree.iter(_qn("w:comment")):
                comments.append(
                    {
                        "id": comment.get(_qn("w:id"), ""),
                        "author": comment.get(_qn("w:author"), ""),
                        "date": comment.get(_qn("w:date"), ""),
                        "text": _tc_extract_text(comment),
                    }
                )

    return {"changes": changes, "comments": comments}


# -- Internal reviewer (in-memory, no filesystem) -----------------------------

_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


class _DocxReviewer:
    """Apply find-replace rules as tracked changes entirely in memory.

    All ZIP contents are held in ``self._files: dict[str, bytes]``.
    """

    def __init__(self, files: dict[str, bytes], author: str):
        self._files = files
        self.author = author
        self.date = __import__("datetime").datetime.now(
            __import__("datetime").timezone.utc
        ).strftime("%Y-%m-%dT%H:%M:%SZ")
        self.comment_id_counter = 0
        self.comments: list[dict] = []
        self.revision_id_counter = 100

        self.doc_root = etree.fromstring(self._files["word/document.xml"])
        self._init_comment_ids()

    # -- id counters -----------------------------------------------------------

    def _init_comment_ids(self):
        if "word/comments.xml" in self._files:
            croot = etree.fromstring(self._files["word/comments.xml"])
            for c in croot.iter(_qn("w:comment")):
                cid = int(c.get(_qn("w:id"), "0"))
                if cid >= self.comment_id_counter:
                    self.comment_id_counter = cid + 1

    def _next_rid(self) -> str:
        self.revision_id_counter += 1
        return str(self.revision_id_counter)

    def _next_comment_id(self) -> int:
        cid = self.comment_id_counter
        self.comment_id_counter += 1
        return cid

    # -- rule application ------------------------------------------------------

    def apply_rules(self, rules: list[dict]) -> int:
        count = 0
        for rule in rules:
            n = self._apply_one_rule(
                rule["find"], rule["replace"], rule.get("comment")
            )
            count += n
        return count

    def _apply_one_rule(self, find: str, replace: str, comment: str | None) -> int:
        body = self.doc_root.find(_qn("w:body"))
        if body is None:
            return 0
        count = 0
        for para in body.iter(_qn("w:p")):
            while True:
                result = self._find_in_paragraph(para, find)
                if result is None:
                    break
                self._replace_in_paragraph(para, result, find, replace, comment)
                count += 1
        return count

    # -- cross-run text search -------------------------------------------------

    def _find_in_paragraph(self, para, find_text: str):
        runs = list(para.iter(_qn("w:r")))
        if not runs:
            return None

        active_runs = []
        for r in runs:
            parent = r.getparent()
            if parent is not None and parent.tag in (_qn("w:del"), _qn("w:ins")):
                continue
            active_runs.append(r)
        if not active_runs:
            return None

        run_texts = []
        for r in active_runs:
            t_elem = r.find(_qn("w:t"))
            run_texts.append(
                t_elem.text if t_elem is not None and t_elem.text else ""
            )

        full_text = "".join(run_texts)
        idx = full_text.find(find_text)
        if idx == -1:
            return None

        start_pos, end_pos = idx, idx + len(find_text)
        cumulative = 0
        start_run_idx = end_run_idx = None
        start_offset = end_offset = 0

        for i, text in enumerate(run_texts):
            run_start = cumulative
            run_end = cumulative + len(text)
            if start_run_idx is None and run_end > start_pos:
                start_run_idx = i
                start_offset = start_pos - run_start
            if run_end >= end_pos:
                end_run_idx = i
                end_offset = end_pos - run_start
                break
            cumulative = run_end

        if start_run_idx is None or end_run_idx is None:
            return None

        return {
            "runs": active_runs[start_run_idx : end_run_idx + 1],
            "start_offset": start_offset,
            "end_offset": end_offset,
        }

    # -- replacement with del/ins markup ---------------------------------------

    def _replace_in_paragraph(self, para, match, find_text, replace_text, comment_text):
        runs = match["runs"]
        start_offset = match["start_offset"]
        end_offset = match["end_offset"]

        rpr_template = runs[0].find(_qn("w:rPr"))
        if rpr_template is not None:
            rpr_template = copy.deepcopy(rpr_template)

        first_t = runs[0].find(_qn("w:t"))
        first_text = first_t.text if first_t is not None and first_t.text else ""
        prefix_text = first_text[:start_offset]

        last_t = runs[-1].find(_qn("w:t"))
        last_text = last_t.text if last_t is not None and last_t.text else ""
        suffix_text = last_text[end_offset:]

        parent = runs[0].getparent()
        insert_pos = list(parent).index(runs[0])
        for r in runs:
            r.getparent().remove(r)

        nodes: list = []

        if prefix_text:
            nodes.append(self._make_run(prefix_text, rpr_template))

        comment_id = None
        if comment_text:
            comment_id = self._next_comment_id()
            cs = etree.Element(_qn("w:commentRangeStart"))
            cs.set(_qn("w:id"), str(comment_id))
            nodes.append(cs)

        # <w:del>
        rid = self._next_rid()
        del_node = etree.Element(_qn("w:del"))
        del_node.set(_qn("w:id"), rid)
        del_node.set(_qn("w:author"), self.author)
        del_node.set(_qn("w:date"), self.date)
        del_node.append(self._make_del_run(find_text, rpr_template))
        nodes.append(del_node)

        # <w:ins>
        ins_node = etree.Element(_qn("w:ins"))
        ins_node.set(_qn("w:id"), self._next_rid())
        ins_node.set(_qn("w:author"), self.author)
        ins_node.set(_qn("w:date"), self.date)
        ins_node.append(self._make_run(replace_text, rpr_template))
        nodes.append(ins_node)

        if comment_text and comment_id is not None:
            ce = etree.Element(_qn("w:commentRangeEnd"))
            ce.set(_qn("w:id"), str(comment_id))
            nodes.append(ce)

            ref_run = etree.Element(_qn("w:r"))
            ref_rpr = etree.SubElement(ref_run, _qn("w:rPr"))
            ref_style = etree.SubElement(ref_rpr, _qn("w:rStyle"))
            ref_style.set(_qn("w:val"), "CommentReference")
            ref_elem = etree.SubElement(ref_run, _qn("w:commentReference"))
            ref_elem.set(_qn("w:id"), str(comment_id))
            nodes.append(ref_run)

            self.comments.append(
                {
                    "id": comment_id,
                    "author": self.author,
                    "date": self.date,
                    "text": comment_text,
                }
            )

        if suffix_text:
            nodes.append(self._make_run(suffix_text, rpr_template))

        for i, node in enumerate(nodes):
            parent.insert(insert_pos + i, node)

    # -- run builders ----------------------------------------------------------

    def _make_run(self, text: str, rpr=None) -> etree._Element:
        run = etree.Element(_qn("w:r"))
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
        t = etree.SubElement(run, _qn("w:t"))
        t.text = text
        t.set(_XML_SPACE, "preserve")
        return run

    def _make_del_run(self, text: str, rpr=None) -> etree._Element:
        run = etree.Element(_qn("w:r"))
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
        dt = etree.SubElement(run, _qn("w:delText"))
        dt.text = text
        dt.set(_XML_SPACE, "preserve")
        return run

    # -- comments & packaging --------------------------------------------------

    def _write_comments(self):
        """Serialize comments into ``word/comments.xml`` in ``self._files``."""
        if not self.comments:
            return

        if "word/comments.xml" in self._files:
            croot = etree.fromstring(self._files["word/comments.xml"])
        else:
            croot = etree.Element(
                _qn("w:comments"),
                nsmap={"w": _NSMAP["w"], "r": _NSMAP["r"]},
            )

        for c in self.comments:
            ce = etree.SubElement(croot, _qn("w:comment"))
            ce.set(_qn("w:id"), str(c["id"]))
            ce.set(_qn("w:author"), c["author"])
            ce.set(_qn("w:date"), c["date"])
            ce.set(_qn("w:initials"), c["author"][:2])

            p = etree.SubElement(ce, _qn("w:p"))
            etree.SubElement(p, _qn("w:pPr"))
            r = etree.SubElement(p, _qn("w:r"))
            rpr = etree.SubElement(r, _qn("w:rPr"))
            rs = etree.SubElement(rpr, _qn("w:rStyle"))
            rs.set(_qn("w:val"), "CommentReference")
            etree.SubElement(r, _qn("w:annotationRef"))

            r2 = etree.SubElement(p, _qn("w:r"))
            t = etree.SubElement(r2, _qn("w:t"))
            t.text = c["text"]
            t.set(_XML_SPACE, "preserve")

        self._files["word/comments.xml"] = etree.tostring(
            croot, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        self._ensure_content_type()
        self._ensure_rels()

    def _ensure_content_type(self):
        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        ct_root = etree.fromstring(self._files["[Content_Types].xml"])
        part_name = "/word/comments.xml"
        ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        for o in ct_root.iter(f"{{{ct_ns}}}Override"):
            if o.get("PartName") == part_name:
                return
        override = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
        override.set("PartName", part_name)
        override.set("ContentType", ct)
        self._files["[Content_Types].xml"] = etree.tostring(
            ct_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    def _ensure_rels(self):
        rels_key = "word/_rels/document.xml.rels"
        rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        rels_root = etree.fromstring(self._files[rels_key])
        for rel in rels_root.iter(f"{{{rels_ns}}}Relationship"):
            if rel.get("Type") == _REL_COMMENTS:
                return
        max_id = 0
        for rel in rels_root.iter(f"{{{rels_ns}}}Relationship"):
            m = re.search(r"(\d+)", rel.get("Id", "rId0"))
            if m:
                max_id = max(max_id, int(m.group(1)))
        new_rel = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
        new_rel.set("Id", f"rId{max_id + 1}")
        new_rel.set("Type", _REL_COMMENTS)
        new_rel.set("Target", "comments.xml")
        self._files[rels_key] = etree.tostring(
            rels_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    def to_bytes(self) -> bytes:
        """Serialize document.xml, comments, and repack into ZIP bytes."""
        self._files["word/document.xml"] = etree.tostring(
            self.doc_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        self._write_comments()

        buf = BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in self._files.items():
                zf.writestr(name, data)
        return buf.getvalue()


# -- Public API: apply_review -------------------------------------------------


def apply_review(
    doc_bytes: bytes,
    rules: list[dict],
    *,
    author: str = "DocKit",
) -> ReviewResult:
    """Apply find-replace rules as tracked changes with optional comments.

    Args:
        doc_bytes: Raw bytes of a .docx file.
        rules: List of dicts, each with ``'find'``, ``'replace'``, and
            optional ``'comment'``.
        author: Author name for the track changes.

    Returns:
        ReviewResult with modified document bytes and replacement count.
    """
    # Load all ZIP entries into memory
    files: dict[str, bytes] = {}
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        for info in zf.infolist():
            files[info.filename] = zf.read(info.filename)

    reviewer = _DocxReviewer(files, author=author)
    count = reviewer.apply_rules(rules)
    return ReviewResult(data=reviewer.to_bytes(), count=count)


# == Markdown to Word ==========================================================

import contextlib

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _md_parse_table_row(line: str) -> list[str]:
    """Parse a Markdown table row, returning cell contents."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _md_is_separator_row(line: str) -> bool:
    """Check whether *line* is a table separator row like ``|---|---|``."""
    line = line.strip()
    if not line.startswith("|"):
        return False
    content = line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return len(content) == 0


def _md_clean_text(text: str) -> str:
    """Strip inline Markdown formatting (bold, italic, code, math)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # *italic*
    text = re.sub(r"`(.+?)`", r"\1", text)  # `code`
    text = re.sub(r"\$(.+?)\$", r"\1", text)  # $math$
    return text


def _md_parse_list_item(line: str) -> tuple[int, str]:
    """Parse a list item, returning ``(indent_level, content)``."""
    stripped = line.lstrip()
    indent = len(line) - len(stripped)
    indent_level = indent // 2  # every 2 spaces = 1 level

    if (
        stripped.startswith("- ")
        or stripped.startswith("* ")
        and not stripped.startswith("**")
        or stripped.startswith("> ")
    ):
        content = stripped[2:]
    elif re.match(r"^\d+\.\s", stripped):
        content = re.sub(r"^\d+\.\s", "", stripped)
    else:
        content = stripped

    return indent_level, content.strip()


def _md_merge_list_items(items: list[str]) -> str:
    """Merge list items into a single paragraph string.

    If the previous item ends with a colon, items are joined directly;
    otherwise a Chinese semicolon is used as separator.
    """
    if not items:
        return ""
    if len(items) == 1:
        return items[0]

    result = items[0]
    for item in items[1:]:
        if result.endswith("\uff1a") or result.endswith(":"):
            result += item
        else:
            result += "\uff1b" + item

    return result


def _md_parse(md_content: str) -> list[dict]:
    """Parse Markdown text into a list of element dicts.

    Recognised element types: ``heading``, ``paragraph``, ``table``,
    ``table_title``, ``figure_title``.
    """
    elements: list[dict] = []
    lines = md_content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # blank line — skip
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            i += 1
            continue

        # heading
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # table title (Chinese convention: starts with 表 + digit)
        if re.match(r"^表\d+", stripped):
            elements.append({"type": "table_title", "text": stripped})
            i += 1
            continue

        # figure title (Chinese convention: starts with 图 + digit)
        if re.match(r"^图\d+", stripped):
            elements.append({"type": "figure_title", "text": stripped})
            i += 1
            continue

        # table (lines starting with |)
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                headers = _md_parse_table_row(table_lines[0])
                rows: list[list[str]] = []
                for tl in table_lines[2:]:
                    if not _md_is_separator_row(tl):
                        rows.append(_md_parse_table_row(tl))
                elements.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # list item (- / * / > / 1.)
        is_list_start = (
            stripped.startswith("- ")
            or (stripped.startswith("* ") and not stripped.startswith("**"))
            or stripped.startswith("> ")
            or bool(re.match(r"^\d+\.\s", stripped))
        )

        if is_list_start:
            list_items: list[tuple[int, str]] = []
            while i < len(lines):
                current = lines[i]
                current_stripped = current.strip()

                if not current_stripped:
                    break

                is_list_item = (
                    current_stripped.startswith("- ")
                    or (
                        current_stripped.startswith("* ")
                        and not current_stripped.startswith("**")
                    )
                    or current_stripped.startswith("> ")
                    or bool(re.match(r"^\d+\.\s", current_stripped))
                    or (current.startswith("  ") and list_items)
                )

                if not is_list_item:
                    break

                il, content = _md_parse_list_item(current)
                if content:
                    list_items.append((il, content))
                i += 1

            if list_items:
                current_para: list[str] = []
                for il, content in list_items:
                    clean_content = _md_clean_text(content)
                    if il == 0:
                        if current_para:
                            elements.append(
                                {
                                    "type": "paragraph",
                                    "text": _md_merge_list_items(current_para),
                                }
                            )
                        current_para = [clean_content]
                    else:
                        current_para.append(clean_content)
                if current_para:
                    elements.append(
                        {
                            "type": "paragraph",
                            "text": _md_merge_list_items(current_para),
                        }
                    )
            continue

        # plain paragraph — collect continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            next_stripped = next_line.strip()

            if not next_stripped:
                break

            if (
                re.match(r"^#{1,4}\s", next_line)
                or next_stripped.startswith("|")
                or next_stripped.startswith("- ")
                or (
                    next_stripped.startswith("* ")
                    and not next_stripped.startswith("**")
                )
                or next_stripped.startswith("> ")
                or re.match(r"^\d+\.\s", next_stripped)
                or next_stripped == "---"
            ):
                break

            para_lines.append(next_line)
            i += 1

        text = " ".join(para_lines)
        text = _md_clean_text(text)
        if text.strip():
            elements.append({"type": "paragraph", "text": text})

    return elements


def _md_set_table_border(
    table, border_color: str = "000000", border_size: int = 4
):
    """Set all borders on a python-docx Table object.

    Args:
        table: python-docx Table object.
        border_color: Hex colour string (default black).
        border_size: Border width in 1/8 pt units (4 = 0.5 pt).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:color"), border_color)
        border.set(qn("w:space"), "0")
        tblBorders.append(border)

    old_borders = tblPr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tblPr.remove(old_borders)

    tblPr.append(tblBorders)


def _md_inject_template_styles(
    doc: Document, template_bytes: bytes
) -> dict:
    """Extract heading styles from *template_bytes* and inject into *doc*.

    Returns a dict with ``doc`` (reloaded Document) and ``available`` (set of
    style names).
    """
    w_ns = _NSMAP["w"]

    # Parse styles.xml from the template
    with zipfile.ZipFile(BytesIO(template_bytes), "r") as zf:
        template_styles_xml = zf.read("word/styles.xml")
    template_root = etree.fromstring(template_styles_xml)

    # Identify target style IDs we want to extract
    target_ids = {
        "a": "Normal",
        "10": "Heading 1",
        "2": "Heading 2",
        "3": "Heading 3",
        "4": "Heading 4",
        "ZDWP": "ZDWP\u6b63\u6587",
        "ZDWP1": "ZDWP\u8868\u540d",
        "ZDWP3": "ZDWP\u8868\u683c\u5185\u5bb9",
        "ZDWP4": "ZDWP\u56fe\u540d",
    }

    style_deps = {
        "10": ["2"],
        "2": ["a"],
        "3": ["4"],
        "4": ["a"],
        "ZDWP": ["a"],
        "ZDWP1": ["a"],
        "ZDWP3": ["a"],
        "ZDWP4": ["a"],
    }

    collected: dict[str, etree._Element] = {}
    for style in template_root.findall(".//w:style", _NSMAP):
        style_id = style.get(f"{{{w_ns}}}styleId")
        style_type = style.get(f"{{{w_ns}}}type")
        if style_type == "paragraph" and style_id in target_ids:
            collected[style_id] = copy.deepcopy(style)

    # Extract docDefaults
    doc_defaults = template_root.find(".//w:docDefaults", _NSMAP)

    # Build a minimal styles XML tree
    new_root = etree.Element(f"{{{w_ns}}}styles", nsmap=_NSMAP)
    if doc_defaults is not None:
        new_root.append(copy.deepcopy(doc_defaults))

    added: set[str] = set()

    def _add(sid: str):
        if sid in added or sid not in collected:
            return
        for dep in style_deps.get(sid, []):
            _add(dep)
        new_root.append(collected[sid])
        added.add(sid)

    for sid in collected:
        _add(sid)

    # Merge into the document's styles.xml via ZIP round-trip
    buf_in = BytesIO()
    doc.save(buf_in)
    buf_in.seek(0)

    files: dict[str, bytes] = {}
    with zipfile.ZipFile(buf_in, "r") as zf:
        for info in zf.infolist():
            files[info.filename] = zf.read(info.filename)

    orig_styles_root = etree.fromstring(files["word/styles.xml"])

    # Remove existing styles that conflict with the template styles
    incoming_ids = {
        s.get(f"{{{w_ns}}}styleId")
        for s in new_root.findall(".//w:style", _NSMAP)
    }
    for s in orig_styles_root.findall(".//w:style", _NSMAP):
        if s.get(f"{{{w_ns}}}styleId") in incoming_ids:
            orig_styles_root.remove(s)

    # Append template styles
    for s in new_root.findall(".//w:style", _NSMAP):
        orig_styles_root.append(copy.deepcopy(s))

    # Update docDefaults
    new_dd = new_root.find(".//w:docDefaults", _NSMAP)
    if new_dd is not None:
        old_dd = orig_styles_root.find(".//w:docDefaults", _NSMAP)
        if old_dd is not None:
            orig_styles_root.remove(old_dd)
        orig_styles_root.insert(0, copy.deepcopy(new_dd))

    files["word/styles.xml"] = etree.tostring(
        orig_styles_root, xml_declaration=True, encoding="UTF-8"
    )

    # Rebuild the docx and reload the Document
    buf_out = BytesIO()
    with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)

    buf_out.seek(0)
    reloaded = Document(buf_out)

    available = {s.name for s in reloaded.styles}
    return {
        "doc": reloaded,
        "available": available,
    }


def _md_write_elements(
    doc: Document, elements: list[dict], available_styles: set[str]
):
    """Write parsed Markdown elements into a python-docx Document."""
    # Determine body / table / figure styles with fallbacks
    body_style = "Normal"
    for candidate in ("ZDWP\u6b63\u6587", "Normal"):
        if candidate in available_styles:
            body_style = candidate
            break

    table_cell_style = (
        "ZDWP\u8868\u683c\u5185\u5bb9"
        if "ZDWP\u8868\u683c\u5185\u5bb9" in available_styles
        else body_style
    )
    table_title_style = (
        "ZDWP\u8868\u540d"
        if "ZDWP\u8868\u540d" in available_styles
        else body_style
    )
    figure_title_style = (
        "ZDWP\u56fe\u540d"
        if "ZDWP\u56fe\u540d" in available_styles
        else body_style
    )

    heading_styles = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
    }

    for elem in elements:
        etype = elem["type"]

        if etype == "heading":
            level = elem["level"]
            style_name = heading_styles.get(level, "Heading 4")
            try:
                doc.add_paragraph(elem["text"], style=style_name)
            except KeyError:
                doc.add_heading(elem["text"], level=level)

        elif etype == "paragraph":
            try:
                doc.add_paragraph(elem["text"], style=body_style)
            except KeyError:
                doc.add_paragraph(elem["text"])

        elif etype == "table_title":
            try:
                doc.add_paragraph(elem["text"], style=table_title_style)
            except KeyError:
                doc.add_paragraph(elem["text"])

        elif etype == "figure_title":
            try:
                doc.add_paragraph(elem["text"], style=figure_title_style)
            except KeyError:
                doc.add_paragraph(elem["text"])

        elif etype == "table":
            headers = elem["headers"]
            data_rows = elem["rows"]
            num_cols = len(headers)
            num_rows = 1 + len(data_rows)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            _md_set_table_border(table)

            # Header row
            header_row = table.rows[0]
            for j, header in enumerate(headers):
                cell = header_row.cells[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = ""
                para = cell.paragraphs[0]
                para.text = header
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                with contextlib.suppress(KeyError):
                    para.style = table_cell_style

            # Data rows
            for ri, row_data in enumerate(data_rows):
                row = table.rows[ri + 1]
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        cell = row.cells[j]
                        cell.vertical_alignment = (
                            WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        )
                        cell.text = ""
                        para = cell.paragraphs[0]
                        para.text = cell_text
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        with contextlib.suppress(KeyError):
                            para.style = table_cell_style


def md_to_docx(
    md_text: str,
    *,
    template_bytes: bytes | None = None,
) -> bytes:
    """Convert Markdown text to a Word document.

    Args:
        md_text: Markdown content string.
        template_bytes: Optional .docx template bytes to extract heading styles
            from.  If None, uses python-docx default styles.

    Returns:
        Word document as bytes.
    """
    elements = _md_parse(md_text)

    if template_bytes is not None:
        doc = Document()
        result = _md_inject_template_styles(doc, template_bytes)
        doc = result["doc"]
        available_styles = result["available"]
    else:
        doc = Document()
        available_styles = {s.name for s in doc.styles}

    _md_write_elements(doc, elements, available_styles)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# == Format checking ===========================================================


@dataclass
class FormatSnapshot:
    """Complete format snapshot of a Word document."""

    page_setup: list[dict] = field(default_factory=list)
    headers_footers: list[dict] = field(default_factory=list)
    watermark: str | None = None
    styles: list[dict] = field(default_factory=list)
    style_usage: dict[str, int] = field(default_factory=dict)
    direct_overrides_count: int = 0
    images_count: int = 0


@dataclass
class CompareResult:
    """Result of comparing two format snapshots."""

    report: str  # Markdown report
    all_ok: bool  # True if no unexpected changes


# -- Constants -----------------------------------------------------------------

_EXPECTED_CHANGES = {"word/document.xml", "word/comments.xml"}
_SAFE_SIDE_EFFECTS = {
    "[Content_Types].xml",
    "word/_rels/document.xml.rels",
    "docProps/core.xml",
    "docProps/app.xml",
    "word/settings.xml",
    "word/endnotes.xml",
    "word/footnotes.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
}

_ALIGN_MAP = {
    "both": "\u4e24\u7aef\u5bf9\u9f50",
    "center": "\u5c45\u4e2d",
    "left": "\u5de6\u5bf9\u9f50",
    "right": "\u53f3\u5bf9\u9f50",
}


# -- Format check helpers ------------------------------------------------------


def _twips_to_cm(val: str) -> float:
    """Convert twips to centimeters."""
    return round(int(val) / 567, 2) if val else 0


def _half_pt(val: str) -> float:
    """Convert Word XML half-point value to points."""
    return int(val) / 2 if val else 0


def _zip_hashes(doc_bytes: bytes) -> dict[str, str]:
    """Compute MD5 hash for each file inside a docx archive."""
    hashes: dict[str, str] = {}
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        for info in zf.infolist():
            data = zf.read(info.filename)
            hashes[info.filename] = hashlib.md5(data).hexdigest()
    return hashes


def _compare_zip_integrity(
    hashes_before: dict[str, str], hashes_after: dict[str, str]
) -> list[dict]:
    """Compare two hash dicts and return a list of differences."""
    diffs: list[dict] = []
    all_keys = set(hashes_before) | set(hashes_after)
    for key in sorted(all_keys):
        h1 = hashes_before.get(key)
        h2 = hashes_after.get(key)
        if h1 == h2:
            continue
        if key in _EXPECTED_CHANGES:
            level = "expected"
        elif key in _SAFE_SIDE_EFFECTS:
            level = "safe"
        else:
            level = "unexpected"

        if h1 is None:
            diffs.append({"file": key, "type": "added", "level": level})
        elif h2 is None:
            diffs.append({"file": key, "type": "removed", "level": level})
        else:
            diffs.append({"file": key, "type": "changed", "level": level})
    return diffs


def _extract_hf_text_raw(tree) -> str:
    """Extract text from a header/footer XML tree, handling PAGE fields and tabs."""
    parts: list[str] = []
    for para in tree.iter(_wqn("p")):
        run_texts: list[str] = []
        in_field = False
        field_instr = ""
        for elem in para.iter():
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "fldChar":
                ftype = elem.get(_wqn("fldCharType"), "")
                if ftype == "begin":
                    in_field = True
                    field_instr = ""
                elif ftype == "end":
                    if "PAGE" in field_instr.upper():
                        run_texts.append("{\u9875\u7801}")
                    in_field = False
            elif tag == "instrText" and in_field:
                field_instr += elem.text or ""
            elif tag == "t" and not in_field:
                run_texts.append(elem.text or "")
            elif tag == "tab":
                run_texts.append(" | ")
        if run_texts:
            line = "".join(run_texts).strip()
            line = re.sub(r"\s{2,}", " ", line)
            if line:
                parts.append(line)
    return " / ".join(parts) if parts else ""


def _extract_snapshot_from_bytes(
    doc_bytes: bytes,
) -> tuple[FormatSnapshot, dict[str, str]]:
    """Extract a FormatSnapshot and zip hashes from document bytes.

    Returns (snapshot, zip_hashes) -- zip_hashes kept separate for comparison.
    """
    snap = FormatSnapshot()
    zhashes = _zip_hashes(doc_bytes)

    _PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    _OREL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as zf:
        names = zf.namelist()

        # -- Relationship mapping: rId -> file path --
        rid_map: dict[str, str] = {}
        rels_path = "word/_rels/document.xml.rels"
        if rels_path in names:
            rels_xml = zf.read(rels_path)
            rels_tree = etree.fromstring(rels_xml)
            for rel in rels_tree.findall(f"{{{_PKG_REL_NS}}}Relationship"):
                rid = rel.get("Id", "")
                target = rel.get("Target", "")
                if not target.startswith("/"):
                    target = "word/" + target
                rid_map[rid] = target

        # -- Pre-extract header/footer text and watermarks --
        hf_text_cache: dict[str, str] = {}
        for name in names:
            if (
                "header" in name.lower() or "footer" in name.lower()
            ) and name.endswith(".xml"):
                xml = zf.read(name)
                htree = etree.fromstring(xml)
                hf_text_cache[name] = _extract_hf_text_raw(htree) or "(\u7a7a)"
                content = xml.decode("utf-8", errors="ignore")
                if (
                    "v:shape" in content
                    or "mso-position-horizontal:center" in content
                ):
                    wm_matches = re.findall(r'string="([^"]+)"', content)
                    if wm_matches:
                        snap.watermark = wm_matches[0]

        # -- Page setup + headers/footers per section --
        doc_xml = zf.read("word/document.xml")
        tree = etree.fromstring(doc_xml)

        section_idx = 0
        for sectPr in tree.iter(_wqn("sectPr")):
            section_idx += 1
            sec: dict = {"section": section_idx}

            pgSz = sectPr.find(_wqn("pgSz"))
            if pgSz is not None:
                w = pgSz.get(_wqn("w"), "")
                h = pgSz.get(_wqn("h"), "")
                orient = pgSz.get(_wqn("orient"), "portrait")
                sec["paper_w_cm"] = _twips_to_cm(w)
                sec["paper_h_cm"] = _twips_to_cm(h)
                sec["orientation"] = orient

            pgMar = sectPr.find(_wqn("pgMar"))
            if pgMar is not None:
                sec["margin_top_cm"] = _twips_to_cm(
                    pgMar.get(_wqn("top"), "0")
                )
                sec["margin_bottom_cm"] = _twips_to_cm(
                    pgMar.get(_wqn("bottom"), "0")
                )
                sec["margin_left_cm"] = _twips_to_cm(
                    pgMar.get(_wqn("left"), "0")
                )
                sec["margin_right_cm"] = _twips_to_cm(
                    pgMar.get(_wqn("right"), "0")
                )
                sec["header_cm"] = _twips_to_cm(
                    pgMar.get(_wqn("header"), "0")
                )
                sec["footer_cm"] = _twips_to_cm(
                    pgMar.get(_wqn("footer"), "0")
                )

            # Read headerReference / footerReference from sectPr
            header_text = None
            footer_text = None
            for ref in sectPr.findall(_wqn("headerReference")):
                htype = ref.get(_wqn("type"), "")
                if htype == "default":
                    rid = ref.get(f"{{{_OREL_NS}}}id", "")
                    fname = rid_map.get(rid, "")
                    header_text = hf_text_cache.get(fname, "(\u7a7a)")
            for ref in sectPr.findall(_wqn("footerReference")):
                ftype = ref.get(_wqn("type"), "")
                if ftype == "default":
                    rid = ref.get(f"{{{_OREL_NS}}}id", "")
                    fname = rid_map.get(rid, "")
                    footer_text = hf_text_cache.get(fname, "(\u7a7a)")

            sec["header"] = header_text
            sec["footer"] = footer_text
            snap.page_setup.append(sec)

        # Inheritance: sections without explicit refs inherit from previous
        for i, sec in enumerate(snap.page_setup):
            if i > 0:
                prev = snap.page_setup[i - 1]
                if sec["header"] is None:
                    sec["header"] = prev["header"]
                    sec["header_inherited"] = True
                if sec["footer"] is None:
                    sec["footer"] = prev["footer"]
                    sec["footer_inherited"] = True

        # Build headers_footers list
        for sec in snap.page_setup:
            hf_entry = {
                "section": sec["section"],
                "header": sec.get("header") or "(\u7a7a)",
                "footer": sec.get("footer") or "(\u7a7a)",
                "header_inherited": sec.get("header_inherited", False),
                "footer_inherited": sec.get("footer_inherited", False),
            }
            snap.headers_footers.append(hf_entry)

        # -- Style definitions --
        if "word/styles.xml" in names:
            styles_xml = zf.read("word/styles.xml")
            stree = etree.fromstring(styles_xml)

            style_names_map: dict[str, str] = {}
            for s in stree.findall(f".//{_wqn('style')}"):
                sid = s.get(_wqn("styleId"), "")
                stype = s.get(_wqn("type"), "")
                if stype != "paragraph":
                    continue

                name_elem = s.find(_wqn("name"))
                name = (
                    name_elem.get(_wqn("val"), sid)
                    if name_elem is not None
                    else sid
                )
                style_names_map[sid] = name

                info: dict = {"id": sid, "name": name}

                rpr = s.find(f".//{_wqn('rPr')}")
                if rpr is not None:
                    rf = rpr.find(_wqn("rFonts"))
                    if rf is not None:
                        info["font_cn"] = rf.get(_wqn("eastAsia"), "")
                        info["font_en"] = rf.get(_wqn("ascii"), "")
                    sz = rpr.find(_wqn("sz"))
                    if sz is not None:
                        info["size_pt"] = _half_pt(sz.get(_wqn("val"), ""))
                    if rpr.find(_wqn("b")) is not None:
                        info["bold"] = True

                ppr = s.find(f".//{_wqn('pPr')}")
                if ppr is not None:
                    jc = ppr.find(_wqn("jc"))
                    if jc is not None:
                        info["align"] = jc.get(_wqn("val"), "")
                    sp = ppr.find(_wqn("spacing"))
                    if sp is not None:
                        line_val = sp.get(_wqn("line"), "")
                        if line_val:
                            info["line_spacing"] = int(line_val)
                    ind = ppr.find(_wqn("ind"))
                    if ind is not None:
                        fc = ind.get(_wqn("firstLineChars"), "")
                        if fc:
                            info["indent_first_chars"] = int(fc)

                snap.styles.append(info)

            # -- Style usage statistics --
            usage_counter: Counter = Counter()
            direct_count = 0
            for para in tree.iter(_wqn("p")):
                ppr = para.find(_wqn("pPr"))
                sid = ""
                if ppr is not None:
                    ps = ppr.find(_wqn("pStyle"))
                    if ps is not None:
                        sid = ps.get(_wqn("val"), "")
                sname = style_names_map.get(sid, sid or "Normal")
                usage_counter[sname] += 1

                for run in para.iter(_wqn("r")):
                    parent = run.getparent()
                    if parent is not None and parent.tag in (
                        _wqn("del"),
                        _wqn("ins"),
                    ):
                        continue
                    rpr = run.find(_wqn("rPr"))
                    if rpr is not None:
                        has_font = rpr.find(_wqn("rFonts")) is not None
                        has_size = rpr.find(_wqn("sz")) is not None
                        if has_font or has_size:
                            direct_count += 1

            snap.style_usage = dict(usage_counter.most_common())
            snap.direct_overrides_count = direct_count

        # -- Image count --
        snap.images_count = len(
            [n for n in names if n.startswith("word/media/")]
        )

    return snap, zhashes


def _format_style_row_check(st: dict) -> str:
    """Format a single style definition row for Markdown table."""
    name = st["name"]
    font_cn = st.get("font_cn", "")
    font_en = st.get("font_en", "")
    size = f"{st['size_pt']}pt" if st.get("size_pt") else "-"
    bold = "**\u662f**" if st.get("bold") else "-"
    align = _ALIGN_MAP.get(st.get("align", ""), st.get("align", "-"))
    if st.get("line_spacing"):
        val = st["line_spacing"]
        line_sp = {
            240: "\u5355\u500d",
            300: "1.25\u500d",
            360: "1.5\u500d",
            480: "2\u500d",
        }.get(val, f"{val}twips")
    else:
        line_sp = "-"
    indent = (
        f"{st['indent_first_chars']}\u5b57\u7b26"
        if st.get("indent_first_chars")
        else "-"
    )
    return (
        f"| {name} | {font_cn or '-'} | {font_en or '-'} "
        f"| {size} | {bold} | {align} | {line_sp} | {indent} |"
    )


def _build_compare_report(
    snap1: FormatSnapshot,
    snap2: FormatSnapshot,
    zhashes1: dict[str, str],
    zhashes2: dict[str, str],
) -> CompareResult:
    """Build a comparison report between two snapshots."""
    lines: list[str] = []
    lines.append("# \u683c\u5f0f\u5bf9\u6bd4\u62a5\u544a\n")

    all_ok = True

    # -- Layer A: ZIP integrity --
    lines.append("## A \u5c42\uff1aZIP \u6587\u4ef6\u5b8c\u6574\u6027\n")
    diffs = _compare_zip_integrity(zhashes1, zhashes2)
    if not diffs:
        lines.append(
            "\u2705 \u6240\u6709\u6587\u4ef6\u5b8c\u5168\u4e00\u81f4"
            "\uff08\u65e0\u4efb\u4f55\u6539\u52a8\uff09\n"
        )
    else:
        expected = [d for d in diffs if d["level"] == "expected"]
        safe = [d for d in diffs if d["level"] == "safe"]
        unexpected = [d for d in diffs if d["level"] == "unexpected"]

        if expected:
            lines.append(
                "\u2705 \u9884\u671f\u53d8\u5316"
                "\uff08\u5ba1\u9605\u4fee\u8ba2\u6838\u5fc3\u6587\u4ef6\uff09\uff1a"
            )
            for d in expected:
                lines.append(f"  - `{d['file']}` ({d['type']})")

        if safe:
            lines.append(
                "\n\u2139\ufe0f  \u5b89\u5168\u526f\u4f5c\u7528"
                "\uff08\u6279\u6ce8/\u4fee\u8ba2\u7684\u57fa\u7840\u8bbe\u65bd\u6587\u4ef6"
                "\uff0c\u4e0d\u5f71\u54cd\u683c\u5f0f\uff09\uff1a"
            )
            for d in safe:
                lines.append(f"  - `{d['file']}` ({d['type']})")

        if unexpected:
            all_ok = False
            lines.append(
                "\n\u274c **\u975e\u9884\u671f\u53d8\u5316**"
                "\uff08\u53ef\u80fd\u5f71\u54cd\u683c\u5f0f/\u9875\u9762/\u56fe\u7247\uff09\uff1a"
            )
            for d in unexpected:
                lines.append(f"  - `{d['file']}` ({d['type']})")

        if not unexpected:
            lines.append(
                f"\n  \u5171 {len(expected) + len(safe)} \u4e2a\u6587\u4ef6\u53d8\u5316"
                "\uff0c\u5168\u90e8\u5728\u9884\u671f\u8303\u56f4\u5185\u3002"
            )
    lines.append("")

    # -- Layer B: semantic comparison --
    lines.append("## B \u5c42\uff1a\u683c\u5f0f\u8bed\u4e49\u5bf9\u6bd4\n")

    # Page setup
    ps1 = snap1.page_setup
    ps2 = snap2.page_setup
    if json.dumps(ps1) == json.dumps(ps2):
        lines.append(
            f"\u2705 \u9875\u9762\u8bbe\u7f6e\uff1a\u672a\u53d8\u5316"
            f"\uff08{len(ps1)} \u4e2a\u5206\u8282\u7b26\uff09"
        )
    else:
        all_ok = False
        lines.append("\u274c \u9875\u9762\u8bbe\u7f6e\uff1a\u6709\u53d8\u5316\uff01")
        lines.append(
            f"  \u539f\u59cb {len(ps1)} \u4e2a\u5206\u8282\u7b26"
            f" \u2192 \u4fee\u6539\u540e {len(ps2)} \u4e2a"
        )

    # Headers/footers
    hf1 = [
        (h.get("section"), h.get("header"), h.get("footer"))
        for h in snap1.headers_footers
    ]
    hf2 = [
        (h.get("section"), h.get("header"), h.get("footer"))
        for h in snap2.headers_footers
    ]
    if hf1 == hf2:
        suffix = (
            f"\uff08{len(hf1)} \u8282\uff09" if hf1 else "\uff08\u65e0\uff09"
        )
        lines.append(
            "\u2705 \u9875\u7709\u9875\u811a\uff1a\u672a\u53d8\u5316" + suffix
        )
    else:
        all_ok = False
        lines.append("\u274c \u9875\u7709\u9875\u811a\uff1a\u6709\u53d8\u5316\uff01")

    # Watermark
    wm1 = snap1.watermark
    wm2 = snap2.watermark
    if wm1 == wm2:
        suffix = (
            f'\uff08"{wm1}"\uff09' if wm1 else "\uff08\u65e0\uff09"
        )
        lines.append("\u2705 \u6c34\u5370\uff1a\u672a\u53d8\u5316" + suffix)
    else:
        all_ok = False
        lines.append(f'\u274c \u6c34\u5370\uff1a"{wm1}" \u2192 "{wm2}"')

    # Style count
    s1 = len(snap1.styles)
    s2 = len(snap2.styles)
    if s1 == s2:
        lines.append(
            f"\u2705 \u6837\u5f0f\u5b9a\u4e49\uff1a\u672a\u53d8\u5316"
            f"\uff08{s1} \u4e2a\u6bb5\u843d\u6837\u5f0f\uff09"
        )
    else:
        all_ok = False
        lines.append(
            f"\u274c \u6837\u5f0f\u5b9a\u4e49\uff1a{s1} \u2192 {s2} \u4e2a"
        )

    # Style usage
    u1 = snap1.style_usage
    u2 = snap2.style_usage
    usage_diff: dict[str, tuple[int, int]] = {}
    for k in set(u1) | set(u2):
        v1 = u1.get(k, 0)
        v2 = u2.get(k, 0)
        if v1 != v2:
            usage_diff[k] = (v1, v2)
    if not usage_diff:
        lines.append(
            "\u2705 \u6837\u5f0f\u4f7f\u7528\uff1a\u6bb5\u843d\u5206\u5e03\u5b8c\u5168\u4e00\u81f4"
        )
    else:
        lines.append(
            f"\u2139\ufe0f  \u6837\u5f0f\u4f7f\u7528\uff1a"
            f"{len(usage_diff)} \u4e2a\u6837\u5f0f\u7684\u6bb5\u843d\u6570\u6709\u53d8\u5316"
        )
        for k, (v1, v2) in sorted(usage_diff.items()):
            lines.append(f"  - {k}: {v1} \u2192 {v2}")

    # Direct overrides
    do1 = snap1.direct_overrides_count
    do2 = snap2.direct_overrides_count
    if do1 == do2:
        lines.append(
            f"\u2705 \u76f4\u63a5\u683c\u5f0f\u8986\u76d6\uff1a{do1} \u4e2a run"
            "\uff08\u672a\u53d8\u5316\uff09"
        )
    else:
        diff = do2 - do1
        sign = "+" if diff > 0 else ""
        lines.append(
            f"\u2139\ufe0f  \u76f4\u63a5\u683c\u5f0f\u8986\u76d6\uff1a"
            f"{do1} \u2192 {do2}\uff08{sign}{diff}\uff09"
            "\uff0c\u53ef\u80fd\u56e0\u4fee\u8ba2\u6807\u8bb0\u5bfc\u81f4"
        )

    # Images
    i1 = snap1.images_count
    i2 = snap2.images_count
    if i1 == i2:
        lines.append(
            f"\u2705 \u5d4c\u5165\u56fe\u7247\uff1a{i1} \u5f20"
            "\uff08\u672a\u53d8\u5316\uff09"
        )
    else:
        all_ok = False
        lines.append(
            f"\u274c \u5d4c\u5165\u56fe\u7247\uff1a{i1} \u2192 {i2} \u5f20"
        )

    lines.append("")
    if all_ok:
        lines.append(
            "## \u7ed3\u8bba\uff1a\u2705 \u683c\u5f0f\u5b8c\u6574\u6027\u901a\u8fc7\n"
        )
        lines.append(
            "\u53ea\u6709\u6587\u672c\u5185\u5bb9\u88ab\u4fee\u6539"
            "\uff08\u4fee\u8ba2\u6807\u8bb0\uff09\uff0c"
            "\u6587\u6863\u683c\u5f0f\u3001\u9875\u9762\u8bbe\u7f6e\u3001"
            "\u56fe\u7247\u7b49\u5747\u672a\u53d8\u5316\u3002"
        )
    else:
        lines.append(
            "## \u7ed3\u8bba\uff1a\u274c \u53d1\u73b0\u975e\u9884\u671f\u53d8\u5316\n"
        )
        lines.append(
            "\u8bf7\u68c0\u67e5\u4e0a\u8ff0\u6807\u8bb0\u4e3a \u274c \u7684\u9879\u76ee\u3002"
        )

    return CompareResult(report="\n".join(lines), all_ok=all_ok)


# -- Format check public API ---------------------------------------------------


def check_format(doc_bytes: bytes) -> FormatSnapshot:
    """Extract a complete format snapshot from a Word document.

    Args:
        doc_bytes: Raw bytes of a .docx file.

    Returns:
        FormatSnapshot with page setup, styles, headers/footers, etc.
    """
    snap, _ = _extract_snapshot_from_bytes(doc_bytes)
    return snap


def format_report(snapshot: FormatSnapshot) -> str:
    """Generate a human-readable Markdown report from a snapshot.

    Args:
        snapshot: A FormatSnapshot obtained from check_format().

    Returns:
        Markdown-formatted report string.
    """
    lines: list[str] = []
    lines.append("# \u6587\u6863\u683c\u5f0f\u62a5\u544a\n")

    # -- 1. Page setup --
    lines.append("## 1. \u9875\u9762\u8bbe\u7f6e\n")
    orientations: Counter = Counter(
        s.get("orientation", "portrait") for s in snapshot.page_setup
    )
    lines.append(
        f"\u5171 {len(snapshot.page_setup)} \u4e2a\u5206\u8282\u7b26"
        f"\uff08\u7eb5\u5411 {orientations.get('portrait', 0)} \u4e2a"
        f"\uff0c\u6a2a\u5411 {orientations.get('landscape', 0)} \u4e2a\uff09\n"
    )
    if snapshot.page_setup:
        s0 = snapshot.page_setup[0]
        lines.append("| \u9879\u76ee | \u503c |")
        lines.append("|------|------|")
        lines.append(
            f"| \u7eb8\u5f20 | {s0.get('paper_w_cm', 0)}cm"
            f" \u00d7 {s0.get('paper_h_cm', 0)}cm (A4) |"
        )
        lines.append(
            f"| \u4e0a\u8fb9\u8ddd | {s0.get('margin_top_cm', 0)}cm |"
        )
        lines.append(
            f"| \u4e0b\u8fb9\u8ddd | {s0.get('margin_bottom_cm', 0)}cm |"
        )
        lines.append(
            f"| \u5de6\u8fb9\u8ddd | {s0.get('margin_left_cm', 0)}cm |"
        )
        lines.append(
            f"| \u53f3\u8fb9\u8ddd | {s0.get('margin_right_cm', 0)}cm |"
        )
        lines.append(
            f"| \u9875\u7709\u8ddd | {s0.get('header_cm', 0)}cm |"
        )
        lines.append(
            f"| \u9875\u811a\u8ddd | {s0.get('footer_cm', 0)}cm |"
        )
    lines.append("")

    # -- 2. Headers/footers & watermark --
    lines.append("## 2. \u9875\u7709\u9875\u811a\n")
    if snapshot.watermark:
        lines.append(f"\u6c34\u5370\uff1a**{snapshot.watermark}**\n")
    else:
        lines.append("\u6c34\u5370\uff1a\u65e0\n")

    if snapshot.headers_footers:
        lines.append("| \u8282 | \u9875\u7709 | \u9875\u811a |")
        lines.append("|------|----------|----------|")
        for hf in snapshot.headers_footers:
            sec_num = hf.get("section", "?")
            h_text = hf.get("header") or "(\u65e0)"
            f_text = hf.get("footer") or "(\u65e0)"
            if hf.get("header_inherited") and h_text != "(\u7a7a)":
                h_text = f"\u2191 {h_text}"
            if hf.get("footer_inherited") and f_text != "(\u7a7a)":
                f_text = f"\u2191 {f_text}"
            lines.append(f"| {sec_num} | {h_text} | {f_text} |")
    else:
        lines.append("\u65e0\u9875\u7709\u9875\u811a\u3002")
    lines.append("")

    # -- 3. Used styles --
    used_names = set(snapshot.style_usage.keys())

    used_styles: list[dict] = []
    unused_styles: list[dict] = []
    for st in snapshot.styles:
        if st["name"] in used_names or st["id"] in used_names:
            used_styles.append(st)
        else:
            unused_styles.append(st)

    usage = snapshot.style_usage
    used_styles.sort(
        key=lambda s: -usage.get(s["name"], usage.get(s["id"], 0))
    )

    lines.append("## 3. \u4f7f\u7528\u4e2d\u7684\u6837\u5f0f\n")
    lines.append(
        f"\u5171 {len(used_styles)} \u4e2a\u6837\u5f0f\u6b63\u5728\u4f7f\u7528"
        f"\uff0c{len(unused_styles)} \u4e2a\u5df2\u5b9a\u4e49\u4f46\u672a\u4f7f\u7528\u3002\n"
    )
    lines.append(
        "| \u6837\u5f0f\u540d | \u6bb5\u843d\u6570 | \u4e2d\u6587\u5b57\u4f53"
        " | \u897f\u6587\u5b57\u4f53 | \u5b57\u53f7 | \u52a0\u7c97"
        " | \u5bf9\u9f50 | \u884c\u8ddd | \u9996\u884c\u7f29\u8fdb |"
    )
    lines.append(
        "|--------|--------|----------|----------|"
        "------|------|------|------|----------|"
    )
    for st in used_styles:
        count = usage.get(st["name"], usage.get(st["id"], 0))
        name = st["name"]
        font_cn = st.get("font_cn", "")
        font_en = st.get("font_en", "")
        size = f"{st['size_pt']}pt" if st.get("size_pt") else "-"
        bold = "**\u662f**" if st.get("bold") else "-"
        align = _ALIGN_MAP.get(st.get("align", ""), st.get("align", "-"))
        if st.get("line_spacing"):
            val = st["line_spacing"]
            line_sp = {
                240: "\u5355\u500d",
                300: "1.25\u500d",
                360: "1.5\u500d",
                480: "2\u500d",
            }.get(val, f"{val}twips")
        else:
            line_sp = "-"
        indent = (
            f"{st['indent_first_chars']}\u5b57\u7b26"
            if st.get("indent_first_chars")
            else "-"
        )
        lines.append(
            f"| {name} | {count} | {font_cn or '-'} | {font_en or '-'}"
            f" | {size} | {bold} | {align} | {line_sp} | {indent} |"
        )
    lines.append("")

    # -- 4. Unused styles --
    if unused_styles:
        lines.append("## 4. \u672a\u4f7f\u7528\u7684\u6837\u5f0f\n")
        lines.append(
            "<details>\n<summary>\u5c55\u5f00\u67e5\u770b "
            + str(len(unused_styles))
            + " \u4e2a\u672a\u4f7f\u7528\u6837\u5f0f</summary>\n"
        )
        lines.append(
            "| \u6837\u5f0f\u540d | \u4e2d\u6587\u5b57\u4f53"
            " | \u897f\u6587\u5b57\u4f53 | \u5b57\u53f7 | \u52a0\u7c97"
            " | \u5bf9\u9f50 | \u884c\u8ddd | \u9996\u884c\u7f29\u8fdb |"
        )
        lines.append(
            "|--------|----------|----------|"
            "------|------|------|------|----------|"
        )
        for st in sorted(unused_styles, key=lambda x: x["name"]):
            lines.append(_format_style_row_check(st))
        lines.append("\n</details>")
    lines.append("")

    # -- 5. Format fingerprint --
    lines.append("## 5. \u683c\u5f0f\u6307\u7eb9\n")
    lines.append("| \u9879\u76ee | \u503c |")
    lines.append("|------|------|")
    lines.append(
        f"| \u76f4\u63a5\u683c\u5f0f\u8986\u76d6\uff08run \u7ea7\uff09"
        f" | {snapshot.direct_overrides_count} \u4e2a |"
    )
    lines.append(
        f"| \u5d4c\u5165\u56fe\u7247 | {snapshot.images_count} \u5f20 |"
    )
    lines.append(
        f"| \u6bb5\u843d\u6837\u5f0f\u5b9a\u4e49\u603b\u6570"
        f" | {len(snapshot.styles)} \u4e2a |"
    )
    total_paras = sum(snapshot.style_usage.values())
    lines.append(
        f"| \u6bb5\u843d\u603b\u6570 | {total_paras} \u4e2a |"
    )
    lines.append("")

    return "\n".join(lines)


def compare_format(
    before_bytes: bytes, after_bytes: bytes
) -> CompareResult:
    """Compare two Word documents and return a format comparison report.

    Args:
        before_bytes: Raw bytes of the original .docx file.
        after_bytes: Raw bytes of the modified .docx file.

    Returns:
        CompareResult with Markdown report and all_ok flag.
    """
    snap1, zhashes1 = _extract_snapshot_from_bytes(before_bytes)
    snap2, zhashes2 = _extract_snapshot_from_bytes(after_bytes)
    return _build_compare_report(snap1, snap2, zhashes1, zhashes2)
