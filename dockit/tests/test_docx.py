"""Tests for dockit.docx module."""

from io import BytesIO

from docx import Document

from dockit.docx import format_text


def _make_docx(*paragraphs: str) -> bytes:
    """Create a minimal DOCX with given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_docx_texts(data: bytes) -> list[str]:
    """Extract all paragraph texts from DOCX bytes."""
    doc = Document(BytesIO(data))
    return [p.text for p in doc.paragraphs if p.text]


class TestFormatText:
    def test_punctuation_fix(self):
        data = _make_docx("hello, world")
        result = format_text(data)
        texts = _read_docx_texts(result.data)
        assert any("\uff0c" in t for t in texts)
        assert result.stats["punctuation"] > 0

    def test_unit_fix(self):
        data = _make_docx("\u9762\u79ef100\u5e73\u65b9\u7c73")
        result = format_text(data)
        texts = _read_docx_texts(result.data)
        assert any("m\u00b2" in t for t in texts)
        assert result.stats["units"] > 0

    def test_quote_fix(self):
        data = _make_docx('"\u6d4b\u8bd5"')
        result = format_text(data)
        texts = _read_docx_texts(result.data)
        # After quote splitting, text may be across multiple runs
        # Just verify stats
        assert result.stats["quotes"] == 2

    def test_empty_document(self):
        data = _make_docx()
        result = format_text(data)
        assert result.stats == {"quotes": 0, "punctuation": 0, "units": 0}
        assert len(result.data) > 0

    def test_table_processing(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "hello, world"
        buf = BytesIO()
        doc.save(buf)

        result = format_text(buf.getvalue())
        assert result.stats["punctuation"] > 0

    def test_returns_valid_docx(self):
        data = _make_docx("\u6d4b\u8bd5\u6587\u672c")
        result = format_text(data)
        # Should be a valid docx that can be opened
        doc = Document(BytesIO(result.data))
        assert len(doc.paragraphs) > 0
